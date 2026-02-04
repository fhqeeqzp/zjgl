"""
招标文件Word文档解析器
从Word文档中提取招标信息
"""
import re
from datetime import datetime
from typing import Dict, Any, Optional


class TenderDocParser:
    """招标文件解析器"""
    
    # 中文数字映射
    CHINESE_NUMBERS = {
        '零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
        '五': 5, '六': 6, '七': 7, '八': 8, '九': 9,
        '十': 10, '百': 100, '千': 1000, '万': 10000, '亿': 100000000
    }
    
    def __init__(self):
        self.extracted_data = {
            'tender_code': '',
            'bidding_name': '',
            'tenderer': '',
            'planned_duration': '',
            'bid_bond': 0.0,
            'bid_deadline': None,
            'control_price': 0.0,
        }
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析招标文件
        :param file_path: Word文档路径
        :return: 提取的数据字典
        """
        try:
            # 尝试使用python-docx库
            from docx import Document
            doc = Document(file_path)
            
            # 提取所有段落文本
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # 提取表格内容
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    tables_text.append(' | '.join(row_text))
            
            all_text = '\n'.join(paragraphs + tables_text)
            
            # 提取各项信息
            self._extract_tender_code(all_text, paragraphs, tables_text)
            self._extract_bidding_name(all_text, paragraphs, tables_text)
            self._extract_tenderer(all_text, paragraphs, tables_text)
            self._extract_planned_duration(all_text, paragraphs, tables_text)
            self._extract_bid_bond(all_text, paragraphs, tables_text)
            self._extract_bid_deadline(all_text, paragraphs, tables_text)
            self._extract_control_price(all_text, paragraphs, tables_text)
            
        except ImportError:
            print("警告: 未安装 python-docx 库，无法解析Word文档")
            print("请运行: pip install python-docx")
        except Exception as e:
            print(f"解析Word文档失败: {e}")
        
        return self.extracted_data
    
    def _chinese_to_number(self, chinese_str: str) -> Optional[float]:
        """
        将中文数字转换为阿拉伯数字
        如：二十六 -> 26, 二千六百五十七 -> 2657
        """
        if not chinese_str:
            return None
        
        # 如果是纯数字，直接返回
        if chinese_str.isdigit():
            return float(chinese_str)
        
        result = 0
        temp = 0
        for char in chinese_str:
            if char in self.CHINESE_NUMBERS:
                num = self.CHINESE_NUMBERS[char]
                if num >= 10000:  # 万、亿
                    if temp == 0:
                        temp = 1
                    result += temp * num
                    temp = 0
                elif num >= 10:  # 十、百、千
                    if temp == 0:
                        temp = 1
                    result += temp * num
                    temp = 0
                else:  # 个位数
                    temp = temp * 10 + num if temp > 0 else num
        
        result += temp
        return float(result) if result > 0 else None
    
    def _parse_amount(self, text: str) -> Optional[float]:
        """
        解析金额文本，支持多种格式
        如：26万元、2657万元、26万、2,657万元、¥26570000
        """
        if not text:
            return None
        
        text = text.strip()
        
        # 1. 尝试匹配 "数字+万元" 格式（最常见）
        pattern_wan = r'(\d+(?:,\d{3})*(?:\.\d+)?)\s*万\s*元?'
        match = re.search(pattern_wan, text)
        if match:
            amount_str = match.group(1).replace(',', '')
            try:
                return float(amount_str) * 10000
            except ValueError:
                pass
        
        # 2. 尝试匹配 "数字+元" 格式
        pattern_yuan = r'(\d+(?:,\d{3})*(?:\.\d+)?)\s*元'
        match = re.search(pattern_yuan, text)
        if match:
            amount_str = match.group(1).replace(',', '')
            try:
                return float(amount_str)
            except ValueError:
                pass
        
        # 3. 尝试匹配中文数字+万元
        pattern_chinese = r'([零一二三四五六七八九十百千万亿]+)\s*万\s*元?'
        match = re.search(pattern_chinese, text)
        if match:
            chinese_num = match.group(1)
            num = self._chinese_to_number(chinese_num)
            if num:
                return num * 10000
        
        # 4. 尝试匹配纯数字（可能已经是元）
        pattern_pure = r'(\d+(?:,\d{3})*(?:\.\d+)?)'
        match = re.search(pattern_pure, text)
        if match:
            amount_str = match.group(1).replace(',', '')
            try:
                return float(amount_str)
            except ValueError:
                pass
        
        return None
    
    def _extract_tender_code(self, all_text: str, paragraphs: list, tables_text: list):
        """提取招标编码"""
        # 常见招标编码格式
        patterns = [
            r'招标编号[：:]\s*([A-Za-z0-9\-]+)',
            r'项目编号[：:]\s*([A-Za-z0-9\-]+)',
            r'招标项目编号[：:]\s*([A-Za-z0-9\-]+)',
            r'编号[：:]\s*([A-Za-z0-9\-]{5,})',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, all_text)
            if match:
                self.extracted_data['tender_code'] = match.group(1).strip()
                return
    
    def _extract_bidding_name(self, all_text: str, paragraphs: list, tables_text: list):
        """提取投标名称/招标项目名称"""
        # 首先尝试从标题中提取
        patterns = [
            r'对(.+?)(?:进行)?(?:招标|采购)',
            r'(.+?)(?:项目|工程)(?:招标|采购)公告',
            r'项目名称[：:]\s*(.+)',
            r'招标项目名称[：:]\s*(.+)',
            r'工程名称[：:]\s*(.+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, all_text)
            if match:
                name = match.group(1).strip()
                # 清理名称
                name = re.sub(r'(招标|采购|项目|工程)$', '', name)
                self.extracted_data['bidding_name'] = name
                return
        
        # 如果上述模式未匹配，尝试从表格中提取
        for row_text in tables_text:
            if '项目名称' in row_text or '工程名称' in row_text:
                parts = row_text.split('|')
                for i, part in enumerate(parts):
                    if '项目名称' in part or '工程名称' in part:
                        if i + 1 < len(parts):
                            self.extracted_data['bidding_name'] = parts[i + 1].strip()
                            return
    
    def _extract_tenderer(self, all_text: str, paragraphs: list, tables_text: list):
        """提取招标人"""
        patterns = [
            r'招标人[：:]\s*([\u4e00-\u9fa5]+(?:公司|局|院|中心|部|单位))',
            r'招标单位[：:]\s*([\u4e00-\u9fa5]+(?:公司|局|院|中心|部|单位))',
            r'采购人[：:]\s*([\u4e00-\u9fa5]+(?:公司|局|院|中心|部|单位))',
            r'建设单位[：:]\s*([\u4e00-\u9fa5]+(?:公司|局|院|中心|部|单位))',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, all_text)
            if match:
                self.extracted_data['tenderer'] = match.group(1).strip()
                return
        
        # 从表格中提取
        for row_text in tables_text:
            if '招标人' in row_text or '招标单位' in row_text:
                parts = row_text.split('|')
                for i, part in enumerate(parts):
                    if '招标人' in part or '招标单位' in part:
                        if i + 1 < len(parts):
                            self.extracted_data['tenderer'] = parts[i + 1].strip()
                            return
    
    def _extract_planned_duration(self, all_text: str, paragraphs: list, tables_text: list):
        """提取计划工期"""
        patterns = [
            r'计划工期[：:]\s*(\d+)\s*(日历天|天|工作日)',
            r'工期[：:]\s*(\d+)\s*(日历天|天|工作日)',
            r'计划总工期[：:]\s*(\d+)\s*(日历天|天|工作日)',
            r'施工工期[：:]\s*(\d+)\s*(日历天|天|工作日)',
            r'(\d+)\s*(日历天|天|工作日)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, all_text)
            if match:
                days = match.group(1)
                unit = match.group(2) if len(match.groups()) > 1 else '日历天'
                self.extracted_data['planned_duration'] = f"{days}{unit}"
                return
    
    def _extract_bid_bond(self, all_text: str, paragraphs: list, tables_text: list):
        """提取投标保证金"""
        # 多种匹配模式，从具体到一般
        patterns = [
            # 模式1: "投标保证金的金额：26万元人民币"
            r'投标保证金的?金额?[：:]\s*([^\n]+?万元?[^\n]*)',
            # 模式2: "投标保证金：26万元"
            r'投标保证金[：:]\s*([^\n]+?万元?[^\n]*)',
            # 模式3: "保证金：26万元"
            r'保证金[：:]\s*([^\n]+?万元?[^\n]*)',
            # 模式4: 包含"投标保证金"的行
            r'投标保证金.*?([\d,]+(?:\.\d+)?)\s*万元',
            # 模式5: 标准格式 "¥26万元" 或 "26万元"
            r'[¥￥]?\s*([\d,]+(?:\.\d+)?)\s*万元',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, all_text, re.IGNORECASE)
            for match in matches:
                # 尝试从匹配文本中提取金额
                matched_text = match.group(0)
                amount = self._parse_amount(matched_text)
                if amount and amount > 0:
                    self.extracted_data['bid_bond'] = amount
                    return
    
    def _extract_bid_deadline(self, all_text: str, paragraphs: list, tables_text: list):
        """提取开标日期/投标截止日期"""
        patterns = [
            r'开标时间[：:]\s*(\d{4}[年/-]\d{1,2}[月/-]\d{1,2}[日]?\s*\d{1,2}[：:]\d{1,2})',
            r'投标截止时间[：:]\s*(\d{4}[年/-]\d{1,2}[月/-]\d{1,2}[日]?\s*\d{1,2}[：:]\d{1,2})',
            r'截止日期[：:]\s*(\d{4}[年/-]\d{1,2}[月/-]\d{1,2}[日]?)',
            r'(\d{4}[年/-]\d{1,2}[月/-]\d{1,2}[日]?)\s*开标',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, all_text)
            if match:
                date_str = match.group(1)
                # 标准化日期格式
                date_str = date_str.replace('年', '-').replace('月', '-').replace('日', '').replace('/', '-')
                try:
                    self.extracted_data['bid_deadline'] = datetime.strptime(date_str.strip(), '%Y-%m-%d %H:%M')
                    return
                except ValueError:
                    try:
                        self.extracted_data['bid_deadline'] = datetime.strptime(date_str.strip(), '%Y-%m-%d')
                        return
                    except ValueError:
                        pass
    
    def _extract_control_price(self, all_text: str, paragraphs: list, tables_text: list):
        """提取招标控制价"""
        # 多种匹配模式，支持各种表述方式
        patterns = [
            # 模式1: "招标控制价为人民币2657万元"
            r'招标控制价(?:为)?[：:]?\s*人民币?\s*([^\n]+?万元?[^\n]*)',
            # 模式2: "招标控制价：2657万元"
            r'招标控制价[：:]\s*([^\n]+?万元?[^\n]*)',
            # 模式3: "最高投标限价：2657万元"
            r'最高投标限价[：:]\s*([^\n]+?万元?[^\n]*)',
            # 模式4: "控制价：2657万元"
            r'控制价[：:]\s*([^\n]+?万元?[^\n]*)',
            # 模式5: "投标控制价：2657万元"
            r'投标控制价[：:]\s*([^\n]+?万元?[^\n]*)',
            # 模式6: "限价：2657万元"
            r'限价[：:]\s*([^\n]+?万元?[^\n]*)',
            # 模式7: "最高限价：2657万元"
            r'最高限价[：:]\s*([^\n]+?万元?[^\n]*)',
            # 模式8: 包含"招标控制价"的行
            r'招标控制价.*?([\d,]+(?:\.\d+)?)\s*万元',
            # 模式9: 包含"控制价"的行
            r'控制价.*?([\d,]+(?:\.\d+)?)\s*万元',
            # 模式10: 标准格式
            r'[¥￥]?\s*([\d,]+(?:\.\d+)?)\s*万元',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, all_text, re.IGNORECASE)
            for match in matches:
                matched_text = match.group(0)
                amount = self._parse_amount(matched_text)
                if amount and amount > 10000:  # 控制价通常较大
                    self.extracted_data['control_price'] = amount
                    return


def parse_tender_document(file_path: str) -> Dict[str, Any]:
    """
    解析招标文件的便捷函数
    :param file_path: Word文档路径
    :return: 提取的数据字典
    """
    parser = TenderDocParser()
    return parser.parse(file_path)
